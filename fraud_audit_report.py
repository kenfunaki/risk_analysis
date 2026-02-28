import os
import json
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from openai import OpenAI


# =========================
# 1) 入力データからLLM用サンプル抽出
# =========================
def build_llm_items(df: pd.DataFrame, top_n: int = 30) -> list[dict]:
    """
    analyze_multiple_row の結果dfから LLMに渡す上位アイテム（辞書リスト）を作る。
    ★偏り防止：会社×科目ごとに均等サンプリングして top_n を作る
    """
    df = df.copy()

    # 想定列（無くても落ちない）
    cols = [
        "company", "account", "year", "quarter", "value",
        "fraud_score", "fraud_anomaly",
        "anomaly", "robust_z", "zscore",
        "qoq_diff", "yoy_diff", "qoq_pct", "qoq_pct_rz",
        "trend_break_z", "sign_flip",
        "explanation",
    ]
    cols = [c for c in cols if c in df.columns]

    # 並び順：fraud_score優先、無ければrobust_z絶対値
    if "fraud_score" in df.columns:
        df = df.sort_values("fraud_score", ascending=False)
    elif "robust_z" in df.columns:
        df = df.sort_values(df["robust_z"].abs(), ascending=False)

    # 期間列
    def add_period(row):
        if pd.notna(row.get("year")) and pd.notna(row.get("quarter")):
            return f"{int(row['year'])}-Q{int(row['quarter'])}"
        return None

    if "period" not in df.columns:
        df["period"] = df.apply(add_period, axis=1)

    # NaN→None
    def to_jsonable(v):
        if pd.isna(v):
            return None
        if isinstance(v, (pd.Timestamp,)):
            return v.isoformat()
        if hasattr(v, "item"):
            try:
                return v.item()
            except Exception:
                pass
        return v

    # -----------------------------
    # ★偏り防止：会社×科目ごとに均等抽出
    # -----------------------------
    if {"company", "account"}.issubset(df.columns):
        grp_cols = ["company", "account"]
        n_groups = df[grp_cols].dropna().drop_duplicates().shape[0]
        per_group = max(1, top_n // max(1, n_groups))

        picked = (
            df.groupby(grp_cols, group_keys=False)
              .head(per_group)
        )

        # 足りない分は全体上位で補完
        if len(picked) < top_n:
            remain = df.loc[~df.index.isin(picked.index)]
            picked = pd.concat([picked, remain.head(top_n - len(picked))])

        top = picked.head(top_n)
    else:
        top = df.head(top_n)

    top = top[cols + (["period"] if "period" not in cols else [])]

    items = []
    for _, r in top.iterrows():
        d = {k: to_jsonable(r.get(k)) for k in top.columns}

        # 日本語キーでLLMへ渡す（A）
        jp = {
            "会社": d.get("company"),
            "勘定科目": d.get("account"),
            "期間": d.get("period"),
            "金額": d.get("value"),

            "リスクスコア": d.get("fraud_score"),
            "リスク兆候": d.get("sign_flip"),
            "外れ値の傾向": d.get("robust_z"),
            "過去の傾向からの乖離度合い": d.get("trend_break_z"),

            "前期比_差分": d.get("qoq_diff"),
            "前年同期比_差分": d.get("yoy_diff"),
            "前期比_増減率": d.get("qoq_pct"),
            "前期比_増減率_外れ補正": d.get("qoq_pct_rz"),

            "特徴要約": (
                f"リスクスコア={d.get('fraud_score')}, "
                f"リスク兆候={d.get('sign_flip')}, "
                f"外れ値の傾向={d.get('robust_z')}, "
                f"z={d.get('zscore')}, "
                f"前期比_差分={d.get('qoq_diff')}, 前年同期比_差分={d.get('yoy_diff')}, "
                f"前期比_増減率={d.get('qoq_pct')}, 前期比_増減率_外れ補正={d.get('qoq_pct_rz')}, "
                f"過去の傾向からの乖離度合い={d.get('trend_break_z')}"
            ),
        }
        items.append(jp)

    return items


# =========================
# 2) GPT-5 で不正リスク評価（構造化出力）
# =========================
def assess_fraud_risk_with_gpt5(items: list[dict], model: str = "gpt-5.2") -> dict:
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError(
            "OPENAI_API_KEY is missing. "
            "Set it in the same shell before running Streamlit, or create src/.env and load it."
        )

    client = OpenAI(api_key=api_key)

    schema = """
{
  "overall_summary": "...",
  "top_themes": ["..."],
  "items": [
    {
      "company": "...",
      "account": "...",
      "period": "YYYY-Qn",
      "risk_level": "High|Medium|Low",
      "themes": ["...", "..."],
      "hypotheses": ["...", "..."],
      "evidence": ["...", "..."],
      "first_checks": ["...", "..."]
    }
  ]
}
""".strip()

    rules = """
用語表記ルール（必ず日本語で記載）：
- QoQ は「前期比」
- YoY は「前年同期比」
- fraud_score は「リスクスコア」
- sign_flip は「リスク兆候」
- robust_z は「外れ値の傾向」
- trend_break_z は「過去の傾向からの乖離度合い」
- evidence / first_checks / overall_summary / top_themes 内で指標名を書く場合も、必ず上記の日本語表記を用いること。
- 入力itemsのキー名（会社/勘定科目/期間/特徴要約/前期比/前年同期比 等）をそのまま使い、英語のキー名に戻さないこと。
""".strip()

    prompt = (
        "あなたは公認不正検査士/内部監査責任者の視点で、以下の「統計検知済みアイテム」を不正リスク（Fraud Risk）として評価してください。\n\n"
        "前提：\n"
        "- リスクスコアは優先順位付けのためのヒューリスティックであり、不正の証明ではない\n"
        "- 外れ値の傾向は統計的外れ、前期比/前年同期比は変化指標\n"
        "- リスク兆候は急落→急回復などの反転を示す（期末調整/戻し仕訳のシグナルになり得る）\n"
        "- 必ず断定を避け、監査仮説として記述する\n\n"
        + rules
        + "\n\n"
        "依頼：\n"
        "0) 本分析は複数会社・複数科目を横断した監査リスク評価であり、単一会社に偏らないこと。\n"
        "1) 全体サマリー（不正リスク観点で何が気になるか、3〜6行）※複数会社を言及すること\n"
        "2) 共通テーマ top_themes（最大5つ）\n"
        "3) itemsごと（会社/勘定科目/期間ごと）に以下を必ず出す：\n"
        "   - risk_level: High / Medium / Low\n"
        "   - themes: 1〜3個（そのアイテムの主要リスクテーマ）\n"
        "   - hypotheses: 1〜3個（監査仮説）\n"
        "   - evidence: 与えた特徴量に基づく根拠（2〜4個）\n"
        "   - first_checks: 最初に確認すべき証跡（2〜5個）\n"
        "4) 可能なら横断要因（複数会社/複数科目に共通）を指摘する\n\n"
        "出力は必ずJSONのみ（このスキーマ厳守）：\n"
        + schema
        + "\n\n"
        "入力items:\n"
        + json.dumps(items, ensure_ascii=False)
    )

    resp = client.responses.create(
        model=model,
        input=prompt,
    )

    text = resp.output_text.strip()

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        start = text.find("{")
        end = text.rfind("}")
        if start != -1 and end != -1 and end > start:
            return json.loads(text[start : end + 1])
        raise

# =========================
# 3) Word監査報告書を生成
# =========================
def export_audit_report_word(
    assessment: dict,
    output_path: str = "fraud_audit_report.docx",
    report_title: str = "財務データ 不正リスク分析（AI支援）監査報告書",
):
    doc = Document()

    # タイトル
    h = doc.add_heading(report_title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(f"作成日: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("※本レポートは統計的異常検知結果に基づくAI支援評価であり、不正の断定ではありません。")

    # 1. 全体サマリー
    doc.add_heading("1. 全体サマリー", level=1)
    doc.add_paragraph(assessment.get("overall_summary", ""))

    # 2. 主要リスクテーマ
    doc.add_heading("2. 主要リスクテーマ", level=1)
    themes = assessment.get("top_themes", []) or []
    if themes:
        for t in themes[:5]:
            doc.add_paragraph(str(t), style="List Bullet")
    else:
        doc.add_paragraph("（該当なし）")

    # 3. 個別リスク評価
    doc.add_heading("3. 個別リスク評価", level=1)

    items = assessment.get("items", []) or []
    if not items:
        doc.add_paragraph("（評価対象アイテムがありません）")
    else:
        for it in items:
            company = it.get("company", "")
            account = it.get("account", "")
            period = it.get("period", "")

            doc.add_heading(f"■ {company} / {account} / {period}", level=2)

            # リスクレベル
            risk = it.get("risk_level", "")
            rp = doc.add_paragraph()
            rrun = rp.add_run(f"リスクレベル: {risk}")
            rrun.bold = True

            # リスクテーマ（追加）
            doc.add_paragraph("リスクテーマ", style=None)
            for t in it.get("themes", []) or []:
                doc.add_paragraph(str(t), style="List Bullet")

            # 仮説
            doc.add_paragraph("想定される不正/統制リスク仮説", style=None)
            for h in it.get("hypotheses", []) or []:
                doc.add_paragraph(str(h), style="List Bullet")

            # 根拠
            doc.add_paragraph("根拠（観測された特徴）", style=None)
            for e in it.get("evidence", []) or []:
                doc.add_paragraph(str(e), style="List Bullet")

            # 優先確認事項
            doc.add_paragraph("優先確認事項（監査/調査手続）", style=None)
            for c in it.get("first_checks", []) or []:
                doc.add_paragraph(str(c), style="List Bullet")

            doc.add_paragraph("")  # 空行

    # 4. 参考（注記）
    doc.add_heading("4. 注記", level=1)
    doc.add_paragraph("・Highは即時レビュー推奨、Mediumは詳細レビュー、Lowはモニタリングを想定。")
    doc.add_paragraph("・本評価は入力データの品質（定義変更、欠損、集計単位）に強く依存する。")
    doc.add_paragraph("・次ステップとして、該当期間の仕訳明細・承認ログ・マスタ変更ログ・配賦ルール変更の突合を推奨。")

    doc.save(output_path)
    return output_path


# =========================
# 4) まとめて実行（CSV→GPT-5→Word）
# =========================
def generate_word_audit_report_from_csv(
    csv_path: str,
    output_docx_path: str = "fraud_audit_report.docx",
    top_n: int = 30,
    model: str = "gpt-5.2",
):
    df = pd.read_csv(csv_path)

    # LLM入力は上位だけ（長すぎるとコスト/精度が悪化）
    items = build_llm_items(df, top_n=top_n)

    assessment = assess_fraud_risk_with_gpt5(items=items, model=model)

    out = export_audit_report_word(assessment, output_path=output_docx_path)
    return out


if __name__ == "__main__":
    # 例：分析結果CSVを指定
    in_csv = "analysis_result_20260228_112244.csv"
    out_docx = "fraud_audit_report.docx"

    path = generate_word_audit_report_from_csv(
        csv_path=in_csv,
        output_docx_path=out_docx,
        top_n=30,
        model="gpt-5.2",  # 「gpt-5」を使いたければここを変更
    )

    print(f"Saved: {path}")